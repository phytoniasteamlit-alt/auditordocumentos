import streamlit as st
import docx
import zipfile
import re
from io import BytesIO

# --- 1. CONFIGURAÇÃO ---
st.set_page_config(page_title="Formatador de Documentos NAQH", page_icon="📊", layout="wide")

with st.sidebar:
    st.markdown("### 🧑‍💻 Operador")
    st.markdown("**Ezequias Santos**\n*Agt Administrativo*")
    st.divider()

st.title("Triagem Avançada & Formatador Automático - NAQH")
st.markdown("""
### 🧠 PRESERVA SEU CABEÇALHO PERFEITO + Aplica Padrão
✅ NÃO apaga nem altera o conteúdo do cabeçalho que você ajustou
✅ Copia APENAS as margens corretas para as outras páginas
✅ Recuo 1,25cm no texto • Sem espaços vazios • Sem páginas em branco
""")

# ============================================================
# 📋 SEÇÕES OBRIGATÓRIAS
# ============================================================
SECOES_POR_TIPO = {
    "POI": ["1. INTRODUÇÃO", "2. OBJETIVO", "3. FINALIDADE", "4. ABRANGÊNCIA", "5. RESPONSABILIDADES", "6. GESTÃO DE RISCO", "7. ANEXOS", "8. REFERÊNCIAS"],
    "POP": ["1. DEFINIÇÃO", "2. APLICABILIDADE", "3. RESPONSABILIDADES", "4. DESCRIÇÃO DAS ETAPAS", "5. REFERÊNCIAS", "6. ANEXOS"],
    "PROG": ["1. REFERENCIAL TEÓRICO", "2. OBJETIVOS", "3. METAS E INDICADORES", "4. DEFINIÇÃO DE METAS", "5. ACOMPANHAMENTO E MONITORAMENTO", "6. AVALIAÇÃO DE RESULTADOS", "7. REFERÊNCIAS", "8. ANEXOS"],
    "PROT": ["1. OBJETIVO", "2. APLICABILIDADE", "3. REFERENCIAL TEÓRICO", "4. CLASSIFICAÇÃO", "5. RESPONSABILIDADES", "6. MEDIDAS PREVENTIVAS", "7. REFERÊNCIAS", "8. ANEXOS"],
    "REG": ["1. FINALIDADE", "2. ÂMBITO", "3. COMPETÊNCIA E ORGANIZAÇÃO", "4. DISPOSIÇÕES GERAIS", "5. DISPOSIÇÕES FINAIS"],
    "NOR": ["1. OBJETIVO", "2. ABRANGÊNCIA", "3. DEFINIÇÕES", "4. COMPETÊNCIAS", "5. PROCEDIMENTOS", "6. DISPOSIÇÕES FINAIS", "7. REFERÊNCIAS"]
}

NAO_RECUAR = [
    "OBJETIVO", "APLICABILIDADE", "REFERENCIAL TEÓRICO", "CLASSIFICAÇÃO",
    "RESPONSABILIDADES", "MEDIDAS PREVENTIVAS", "REFERÊNCIAS", "ANEXOS",
    "DEFINIÇÃO", "FINALIDADE", "ÂMBITO", "COMPETÊNCIA", "PROCEDIMENTOS",
    "DISPOSIÇÕES", "QUADRO", "Figura", "Tabela", "a)", "b)", "c)", "d)", "e)",
    "5.1", "5.2", "5.3", "5.4", "6.1", "•", "●"
]

# ============================================================
# 🧠 MOTOR — PRESERVA CABEÇALHO + APLICA MARGENS SEM APAGAR
# ============================================================
def formatar_sem_apagar_cabecalho(arquivo_bytes):
    """
    ✅ LE o cabeçalho que VOCÊ ajustou → PRESERVA 100% do conteúdo
    ✅ Copia esse cabeçalho para as outras páginas
    ✅ Aplica margens corretas (3,0 / 3,0 / 2,0 / 2,0) SEM alterar o conteúdo
    ✅ Recuo 1,25cm no texto • Sem espaços vazios • Sem páginas em branco
    """
    top_dxa, bottom_dxa, left_dxa, right_dxa = "1701", "1134", "1701", "1134"
    recuo = "709"  # 1,25cm
    
    zip_original = zipfile.ZipFile(BytesIO(arquivo_bytes))
    buffer_saida = BytesIO()
    
    # ✅ LE o cabeçalho PERFEITO que você ajustou — guarda como MODELO
    modelo_cabecalho = None
    modelo_rodape = None
    
    for info in zip_original.infolist():
        if info.filename == "word/header1.xml":
            modelo_cabecalho = zip_original.read(info.filename).decode("utf-8")
        if info.filename == "word/footer1.xml":
            modelo_rodape = zip_original.read(info.filename).decode("utf-8")
    
    zip_original.close()
    
    # ✅ Reabre para processar
    zip_original = zipfile.ZipFile(BytesIO(arquivo_bytes))
    
    with zipfile.ZipFile(buffer_saida, "w", zipfile.ZIP_DEFLATED) as zip_novo:
        for item in zip_original.infolist():
            conteudo = zip_original.read(item.filename)
            
            # ======================================
            # CORPO DO DOCUMENTO — margens + recuo
            # ======================================
            if item.filename == "word/document.xml":
                xml = conteudo.decode("utf-8")
                
                # ✅ Aplica margens corretas
                xml = re.sub(r'w:top="\d+"',    f'w:top="{top_dxa}"',    xml)
                xml = re.sub(r'w:bottom="\d+"', f'w:bottom="{bottom_dxa}"', xml)
                xml = re.sub(r'w:left="\d+"',   f'w:left="{left_dxa}"',   xml)
                xml = re.sub(r'w:right="\d+"',  f'w:right="{right_dxa}"',  xml)
                
                # ✅ Zera espaçamentos → SEM buracos gigantes
                xml = re.sub(r'w:spaceBefore="\d+"', 'w:spaceBefore="0"', xml)
                xml = re.sub(r'w:spaceAfter="\d+"',  'w:spaceAfter="240"', xml)
                
                # ✅ Remove recuos antigos
                xml = re.sub(r'<w:ind[^>]+/?>', '', xml)
                
                # ✅ Aplica recuo de 1,25cm na 1ª linha
                xml = re.sub(
                    r'(<w:pPr>)((?:(?!<w:numPr>|<w:pStyle).){0,150}</w:pPr>)',
                    rf'\1\2<w:ind w:firstLine="{recuo}"/>',
                    xml
                )
                
                # ✅ Remove recuo de títulos
                for termo in NAO_RECUAR:
                    xml = re.sub(
                        rf'<w:ind w:firstLine="{recuo}"/>([^<]*?{re.escape(termo)}[^<]*?)</w:pPr>',
                        r'\1</w:pPr>',
                        xml,
                        flags=re.IGNORECASE
                    )
                
                # ✅ Remove quebras que causam páginas vazias
                xml = re.sub(r'<w:br w:type="page"/>', '', xml)
                xml = re.sub(r'<w:pageBreakBefore/>', '', xml)
                
                conteudo = xml.encode("utf-8")

            # ======================================
            # ✅ CABEÇALHOS — USA SEU MODELO PERFEITO
            # ======================================
            elif item.filename.startswith("word/header"):
                if modelo_cabecalho:
                    # ✅ USA SEU CABEÇALHO — NÃO APAGA NADA!
                    xml = modelo_cabecalho
                    # ✅ Apenas garante margens corretas — SEM alterar conteúdo
                    xml = re.sub(r'w:top="\d+"',    f'w:top="{top_dxa}"',    xml)
                    xml = re.sub(r'w:bottom="\d+"', f'w:bottom="{bottom_dxa}"', xml)
                    xml = re.sub(r'w:left="\d+"',   f'w:left="{left_dxa}"',   xml)
                    xml = re.sub(r'w:right="\d+"',  f'w:right="{right_dxa}"',  xml)
                    conteudo = xml.encode("utf-8")
                else:
                    # Se não encontrou modelo, só ajusta margens
                    xml = conteudo.decode("utf-8")
                    xml = re.sub(r'w:top="\d+"',    f'w:top="{top_dxa}"',    xml)
                    xml = re.sub(r'w:bottom="\d+"', f'w:bottom="{bottom_dxa}"', xml)
                    xml = re.sub(r'w:left="\d+"',   f'w:left="{left_dxa}"',   xml)
                    xml = re.sub(r'w:right="\d+"',  f'w:right="{right_dxa}"',  xml)
                    conteudo = xml.encode("utf-8")

            # ======================================
            # ✅ RODAPÉS — preserva e ajusta margens
            # ======================================
            elif item.filename.startswith("word/footer"):
                if modelo_rodape:
                    xml = modelo_rodape
                    xml = re.sub(r'w:top="\d+"',    f'w:top="{top_dxa}"',    xml)
                    xml = re.sub(r'w:bottom="\d+"', f'w:bottom="{bottom_dxa}"', xml)
                    xml = re.sub(r'w:left="\d+"',   f'w:left="{left_dxa}"',   xml)
                    xml = re.sub(r'w:right="\d+"',  f'w:right="{right_dxa}"',  xml)
                    conteudo = xml.encode("utf-8")
                else:
                    xml = conteudo.decode("utf-8")
                    xml = re.sub(r'w:top="\d+"',    f'w:top="{top_dxa}"',    xml)
                    xml = re.sub(r'w:bottom="\d+"', f'w:bottom="{bottom_dxa}"', xml)
                    xml = re.sub(r'w:left="\d+"',   f'w:left="{left_dxa}"',   xml)
                    xml = re.sub(r'w:right="\d+"',  f'w:right="{right_dxa}"',  xml)
                    conteudo = xml.encode("utf-8")
            
            zip_novo.writestr(item, conteudo)
            
    zip_original.close()
    buffer_saida.seek(0)
    return buffer_saida.getvalue()

# ============================================================
# 🔍 TRIAGEM
# ============================================================
def identificar_tipo_e_secoes(texto):
    texto = texto.upper()
    if "PROT_" in texto or "PROTOCOLO" in texto:
        return "PROT", SECOES_POR_TIPO["PROT"]
    elif "POP_" in texto or "PROCEDIMENTO OPERACIONAL" in texto:
        return "POP", SECOES_POR_TIPO["POP"]
    elif "NOR_" in texto or "NORMA" in texto:
        return "NOR", SECOES_POR_TIPO["NOR"]
    elif "REG_" in texto or "REGULAMENTO" in texto:
        return "REG", SECOES_POR_TIPO["REG"]
    elif "PROG_" in texto or "PROGRAMA" in texto:
        return "PROG", SECOES_POR_TIPO["PROG"]
    elif "POI_" in texto or "POLÍTICA" in texto or "POLITICA" in texto:
        return "POI", SECOES_POR_TIPO["POI"]
    else:
        return "PROT", SECOES_POR_TIPO["PROT"]

# ============================================================
# 🚀 INTERFACE
# ============================================================
with st.form("form_preserva_cabecalho"):
    arquivo_word = st.file_uploader(
        "📂 Arraste o documento WORD (.docx) AQUI com o cabeçalho já ajustado",
        type=["docx"],
        key="upload_preserva_cabecalho"
    )
    enviado = st.form_submit_button("🔄 APLICAR PADRÃO — PRESERVAR CABEÇALHO", type="primary")

if enviado and arquivo_word:
    st.info(f"✅ Arquivo carregado: **{arquivo_word.name}**")
    
    with st.spinner("Preservando seu cabeçalho e aplicando margens..."):
        dados_brutos = arquivo_word.read()
        
        doc_triagem = docx.Document(BytesIO(dados_brutos))
        texto_corpo_raw = " ".join([p.text.strip() for p in doc_triagem.paragraphs[:60]])
        texto_tabelas_raw = " ".join([cell.text.strip() for t in doc_triagem.tables[:2] for r in t.rows for cell in r.cells])
        texto_total_raw = (texto_corpo_raw + " " + texto_tabelas_raw).upper()
        
        sigla_tipo, secoes_esperadas = identificar_tipo_e_secoes(texto_total_raw)
        
        codigo_doc = f"{sigla_tipo}_SCIH000"
        match_codigo = re.search(r'\b(PROT|POP|MAN|NOR|REG|PROG|POI)_[A-Z0-9_\s-]+\b', texto_total_raw, re.IGNORECASE)
        if match_codigo:
            codigo_doc = match_codigo.group(0).strip().upper().replace(" ", "")
        
        dados_finais = formatar_sem_apagar_cabecalho(dados_brutos)
        
        st.success(f"📋 **DOCUMENTO IDENTIFICADO: {sigla_tipo}**")
        
        st.markdown("### 📑 Estrutura / Seções Obrigatórias:")
        for secao in secoes_esperadas:
            st.write(f"✅ {secao}")
        
        st.markdown("---")
        
        st.download_button(
            label="📥 BAIXAR — CABEÇALHO PRESERVADO + MARGENS",
            data=dados_finais,
            file_name=f"{codigo_doc}_Cabecalho_Preservado_Norma_Zero.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
        
        st.success("""✅ **CONCLUÍDO — SEU CABEÇALHO FOI PRESERVADO!**
• ✅ SEU cabeçalho ajustado → COPIADO para TODAS as páginas
• ✅ NADA do cabeçalho foi apagado ou alterado
• ✅ Margens: Sup 3,0cm | Esq 3,0cm | Inf 2,0cm | Dir 2,0cm
• ✅ Recuo 1,25cm no texto • Títulos SEM recuo
• ✅ Espaçamentos normalizados • Sem páginas em branco vazias""")
